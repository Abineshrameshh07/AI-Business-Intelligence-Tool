"""
analysis.py
------------
Turns a CSV file into simple, chart-ready statistics.
"""

import re
import sqlite3
from collections import Counter
from dotenv import load_dotenv
import os
import pandas as pd
from docx import Document   

def save_parsed_data(file_path: str, df: pd.DataFrame):
    """
    Saves the parsed DataFrame as a JSON file next to the uploaded file.
    This lets the drilldown endpoint filter rows without re-parsing the file.
    e.g. uploads/sales.csv → uploads/sales.csv.parsed.json
    """
    parsed_path = file_path + ".parsed.json"
    df.to_json(parsed_path, orient="records")

load_dotenv()


def analyze_csv(file_path: str) -> dict:
    df = pd.read_csv(file_path)

    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    categorical_cols = df.select_dtypes(exclude="number").columns.tolist()

    summary = {
        "row_count": len(df),
        "column_count": len(df.columns),
        "columns": df.columns.tolist(),
        "numeric_columns": numeric_cols,
        "categorical_columns": categorical_cols,
    }

    chart_config = build_chart_config(df, numeric_cols, categorical_cols)
    save_parsed_data(file_path, df)
    sample_rows = df.head(5).to_dict(orient="records")
    ai_insights = get_ai_insights(summary, chart_config, sample_rows)

    return {
        "summary": summary,
        "chart_config": chart_config,
        "ai_insights": ai_insights,
    }


def build_chart_config(df: pd.DataFrame, numeric_cols: list, categorical_cols: list) -> dict:
    if categorical_cols and numeric_cols:
        cat_col = categorical_cols[0]
        num_col = numeric_cols[0]
        grouped = df.groupby(cat_col)[num_col].sum().reset_index()
        return {
            "chart_type": "bar",
            "title": f"Total {num_col} by {cat_col}",
            "labels": grouped[cat_col].tolist(),
            "values": grouped[num_col].tolist(),
        }

    if numeric_cols:
        num_col = numeric_cols[0]
        return {
            "chart_type": "line",
            "title": f"{num_col} over rows",
            "labels": [str(i) for i in range(len(df))],
            "values": df[num_col].tolist(),
        }

    if categorical_cols:
        cat_col = categorical_cols[0]
        counts = df[cat_col].value_counts().reset_index()
        counts.columns = [cat_col, "count"]
        return {
            "chart_type": "pie",
            "title": f"Count of {cat_col}",
            "labels": counts[cat_col].tolist(),
            "values": counts["count"].tolist(),
        }

    return {"chart_type": None, "title": "No chartable columns found", "labels": [], "values": []}


STOPWORDS = {
    "the", "and", "a", "an", "of", "to", "in", "is", "it", "for", "on",
    "with", "as", "was", "were", "are", "this", "that", "by", "at", "be",
    "or", "from", "has", "have", "had", "not", "but", "which", "will",
}


def make_column_names_unique(columns: list) -> list:
    seen = {}
    unique_columns = []
    for col in columns:
        if col not in seen:
            seen[col] = 1
            unique_columns.append(col)
        else:
            seen[col] += 1
            unique_columns.append(f"{col}_{seen[col]}")
    return unique_columns


def analyze_docx(file_path: str) -> dict:
    doc = Document(file_path)

    # --- Case 1: document has a table ---
    if doc.tables:
        table = doc.tables[0]
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        header = make_column_names_unique(rows[0])
        data_rows = rows[1:]
        df = pd.DataFrame(data_rows, columns=header)

        for col in df.columns:
            try:
                df[col] = pd.to_numeric(df[col])
            except (ValueError, TypeError):
                pass

        numeric_cols = df.select_dtypes(include="number").columns.tolist()
        categorical_cols = df.select_dtypes(exclude="number").columns.tolist()

        summary = {
            "row_count": len(df),
            "column_count": len(df.columns),
            "columns": df.columns.tolist(),
            "numeric_columns": numeric_cols,
            "categorical_columns": categorical_cols,
            "source": "table found in document",
        }
        chart_config = build_chart_config(df, numeric_cols, categorical_cols)
        save_parsed_data(file_path, df)
        sample_rows = df.head(5).to_dict(orient="records")
        ai_insights = get_ai_insights(summary, chart_config, sample_rows)
        return {
        "upload_id": None,   # gets filled in by main.py
        "summary": summary,
        "chart_config": chart_config,
        "ai_insights": ai_insights,
    }

    # --- Case 2: no table, use word frequency ---
    full_text = " ".join(p.text for p in doc.paragraphs)
    words = re.findall(r"[a-zA-Z]+", full_text.lower())
    meaningful_words = [w for w in words if w not in STOPWORDS and len(w) > 3]
    top_words = Counter(meaningful_words).most_common(10)

    summary = {
        "row_count": len(doc.paragraphs),
        "column_count": None,
        "columns": [],
        "numeric_columns": [],
        "categorical_columns": [],
        "source": "no table found, analyzed word frequency instead",
    }

    if not top_words:
        chart_config = {"chart_type": None, "title": "No text found to analyze", "labels": [], "values": []}
    else:
        labels, values = zip(*top_words)
        chart_config = {
            "chart_type": "bar",
            "title": "Most frequent words in document",
            "labels": list(labels),
            "values": list(values),
        }

    ai_insights = get_ai_insights(summary, chart_config, [])
    return {
        "upload_id": None,   # gets filled in by main.py
        "summary": summary,
        "chart_config": chart_config,
        "ai_insights": ai_insights,
    }


def analyze_sql(file_path: str) -> dict:
    with open(file_path, "r") as f:
        script = f.read()

    statements = [s.strip() for s in script.split(";") if s.strip()]
    conn = sqlite3.connect(":memory:")
    result_df = None

    for stmt in statements:
        if stmt.lower().startswith("select"):
            try:
                result_df = pd.read_sql_query(stmt, conn)
            except Exception as e:
                raise ValueError(f"Error running SELECT statement: {e}")
        else:
            conn.execute(stmt)

    conn.commit()
    conn.close()

    if result_df is None or result_df.empty:
        return {
            "summary": {"row_count": 0, "column_count": 0, "columns": [], "source": "no SELECT result found"},
            "chart_config": {"chart_type": None, "title": "No data returned from SQL file", "labels": [], "values": []},
            "ai_insights": "No data was returned from this SQL file to analyze.",
        }

    numeric_cols = result_df.select_dtypes(include="number").columns.tolist()
    categorical_cols = result_df.select_dtypes(exclude="number").columns.tolist()

    summary = {
        "row_count": len(result_df),
        "column_count": len(result_df.columns),
        "columns": result_df.columns.tolist(),
        "numeric_columns": numeric_cols,
        "categorical_columns": categorical_cols,
        "source": "result of SELECT statement in .sql file",
    }
    chart_config = build_chart_config(result_df, numeric_cols, categorical_cols)
    save_parsed_data(file_path, result_df)
    sample_rows = result_df.head(5).to_dict(orient="records")
    ai_insights = get_ai_insights(summary, chart_config, sample_rows)

    return {
        "upload_id": None,   # gets filled in by main.py
        "summary": summary,
        "chart_config": chart_config,
        "ai_insights": ai_insights,
    }


def get_ai_insights(summary: dict, chart_config: dict, sample_rows: list) -> str:
    from groq import Groq
    client = Groq(api_key=os.getenv("GROQ_API_KEY"))

    prompt = f"""
You are a data analyst assistant. A user has uploaded a data file and I need you to analyze it.

Here is a summary of the data:
- Rows: {summary.get('row_count')}
- Columns: {summary.get('columns')}
- Numeric columns: {summary.get('numeric_columns', [])}
- Categorical columns: {summary.get('categorical_columns', [])}

Sample data (first few rows):
{sample_rows}

A chart has been auto-generated with these settings:
- Chart type: {chart_config.get('chart_type')}
- Title: {chart_config.get('title')}
- Labels: {chart_config.get('labels')}
- Values: {chart_config.get('values')}

Please provide:
1. A short 2-3 sentence summary of what this data is about
2. The most interesting pattern or insight you can see
3. One specific recommendation based on the data

Keep your response concise and friendly, written for a non-technical user.
"""

    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        max_tokens=400,
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content